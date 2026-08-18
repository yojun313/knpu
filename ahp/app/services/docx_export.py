"""설문지 Word(.docx) 내보내기 — 종이로 인쇄해 나눠줄 설문지를 만든다.

PLAN.md 6.2: "같은 계층의 페이지 넘김 자제". python-docx는 "표를 페이지 사이에서
쪼개지 않기"를 공개 API로 노출하지 않아서, 각 행에 OOXML의 w:cantSplit을 직접
심어야 한다 — 짧은 스니펫이지만 없으면 질문 하나가 페이지 중간에서 잘릴 수 있다.
"""

import io

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

SAATY_LEVELS = [
    9,
    8,
    7,
    6,
    5,
    4,
    3,
    2,
    1,
    2,
    3,
    4,
    5,
    6,
    7,
    8,
    9,
]  # 왼쪽부터: item_a 절대적(9) ... 동일(1) ... item_b 절대적(9) — 1~9 전 구간


def _prevent_row_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def _set_cell_text(cell, text, *, bold=False, size=10, align_center=False):
    cell.text = ""
    p = cell.paragraphs[0]
    if align_center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)


def build_survey_docx(survey: dict, nodes_by_uuid: dict) -> io.BytesIO:
    doc = Document()

    title = doc.add_heading(survey.get("title") or "AHP 설문지", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    if survey.get("intro_text"):
        p = doc.add_paragraph(survey["intro_text"])
        p.paragraph_format.space_after = Pt(10)

    if survey.get("consent_text"):
        h = doc.add_paragraph()
        h.add_run("안내 및 동의").bold = True
        doc.add_paragraph(survey["consent_text"])

    doc.add_paragraph()
    meta = doc.add_paragraph()
    meta.add_run(
        "응답자: ______________________    소속/경력: ______________________    일자: ______________________"
    )
    doc.add_paragraph(
        "※ 아래 각 항목 쌍에서 두 항목의 상대적 중요도를 비교해, 해당하는 칸에 표시(○ 또는 ✓)해 주세요. "
        "가운데(1)는 두 항목이 동일하게 중요함을 뜻합니다."
    )

    node_descriptions = survey.get("node_descriptions", {})

    for m in survey.get("matrices", []):
        parent_name = nodes_by_uuid.get(m["parent_uuid"], {}).get("name", "")
        parent_desc = node_descriptions.get(m["parent_uuid"], "")

        h = doc.add_heading(f"'{parent_name}' 측면 비교", level=2)
        h.paragraph_format.keep_with_next = True
        if parent_desc:
            doc.add_paragraph(parent_desc).italic = True
        q = doc.add_paragraph(m.get("question_text", ""))
        q.paragraph_format.keep_with_next = True

        child_uuids = m["child_uuids"]
        for i in range(len(child_uuids)):
            for j in range(i + 1, len(child_uuids)):
                name_a = nodes_by_uuid.get(child_uuids[i], {}).get(
                    "name", child_uuids[i]
                )
                name_b = nodes_by_uuid.get(child_uuids[j], {}).get(
                    "name", child_uuids[j]
                )

                table = doc.add_table(rows=2, cols=1 + len(SAATY_LEVELS) + 1)
                table.style = "Table Grid"
                table.autofit = False
                for row in table.rows:
                    _prevent_row_split(row)

                header = table.rows[0]
                _set_cell_text(header.cells[0], name_a, bold=True, size=9)
                for k, lvl in enumerate(SAATY_LEVELS):
                    label = str(lvl) if k <= len(SAATY_LEVELS) // 2 else f"1/{lvl}"
                    _set_cell_text(
                        header.cells[1 + k], label, align_center=True, size=7
                    )
                _set_cell_text(header.cells[-1], name_b, bold=True, size=9)

                mark_row = table.rows[1]
                for k in range(len(SAATY_LEVELS) + 2):
                    _set_cell_text(mark_row.cells[k], "☐", align_center=True, size=10)

                # 17칸(9~1~1/9)이 되면서 칸마다 폭을 줄여야 기본 여백(1인치, 가용
                # 폭 ~15.9cm A4 기준) 안에 맞는다: 2*1.8 + 17*0.7 = 15.5cm.
                for cell, width in zip(
                    table.columns, [Cm(1.8)] + [Cm(0.7)] * len(SAATY_LEVELS) + [Cm(1.8)]
                ):
                    for c in cell.cells:
                        c.width = width

                doc.add_paragraph().paragraph_format.space_after = Pt(4)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf
